"""
Service de parsing de documents multi-formats.
Supporte: PDF, DOCX, TXT, MD, HTML
"""

import logging
from pathlib import Path
from typing import Dict, List, Optional
from dataclasses import dataclass
from enum import Enum

logger = logging.getLogger(__name__)


class DocumentFormat(str, Enum):
    """Formats de documents supportés"""
    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"
    MARKDOWN = "md"
    HTML = "html"
    UNKNOWN = "unknown"


@dataclass
class ParsedDocument:
    """Document parsé avec métadonnées"""
    content: str
    format: DocumentFormat
    metadata: Dict[str, str]
    page_count: Optional[int] = None
    word_count: Optional[int] = None


class DocumentParserService:
    """Service de parsing de documents multi-formats"""

    def __init__(self):
        self.supported_formats = {
            ".pdf": DocumentFormat.PDF,
            ".docx": DocumentFormat.DOCX,
            ".txt": DocumentFormat.TXT,
            ".md": DocumentFormat.MARKDOWN,
            ".html": DocumentFormat.HTML,
            ".htm": DocumentFormat.HTML,
        }

    def detect_format(self, file_path: str) -> DocumentFormat:
        """Détecte le format d'un document à partir de son extension"""
        suffix = Path(file_path).suffix.lower()
        return self.supported_formats.get(suffix, DocumentFormat.UNKNOWN)

    async def parse_file(self, file_path: str, metadata: Optional[Dict] = None) -> ParsedDocument:
        """Parse un document et retourne son contenu avec métadonnées"""
        doc_format = self.detect_format(file_path)

        if doc_format == DocumentFormat.UNKNOWN:
            raise ValueError(f"Format de document non supporté: {file_path}")

        logger.info(f"[DocumentParser] Parsing {doc_format.value}: {file_path}")

        # Parser selon le format
        if doc_format == DocumentFormat.PDF:
            return await self._parse_pdf(file_path, metadata or {})
        elif doc_format == DocumentFormat.DOCX:
            return await self._parse_docx(file_path, metadata or {})
        elif doc_format == DocumentFormat.TXT:
            return await self._parse_text(file_path, metadata or {})
        elif doc_format == DocumentFormat.MARKDOWN:
            return await self._parse_markdown(file_path, metadata or {})
        elif doc_format == DocumentFormat.HTML:
            return await self._parse_html(file_path, metadata or {})

        raise ValueError(f"Parser non implémenté pour: {doc_format}")

    async def parse_content(
        self,
        content: str,
        format: DocumentFormat,
        metadata: Optional[Dict] = None
    ) -> ParsedDocument:
        """Parse du contenu brut déjà en mémoire"""
        word_count = len(content.split())

        return ParsedDocument(
            content=content,
            format=format,
            metadata=metadata or {},
            word_count=word_count
        )

    async def _parse_pdf(self, file_path: str, metadata: Dict) -> ParsedDocument:
        """Parse un fichier PDF"""
        try:
            from pypdf import PdfReader

            reader = PdfReader(file_path)
            pages = []

            for page in reader.pages:
                text = page.extract_text()
                if text.strip():
                    pages.append(text)

            content = "\n\n".join(pages)

            # Extraire métadonnées PDF
            pdf_metadata = reader.metadata or {}
            metadata.update({
                "title": pdf_metadata.get("/Title", ""),
                "author": pdf_metadata.get("/Author", ""),
                "subject": pdf_metadata.get("/Subject", ""),
                "creator": pdf_metadata.get("/Creator", ""),
            })

            return ParsedDocument(
                content=content,
                format=DocumentFormat.PDF,
                metadata=metadata,
                page_count=len(reader.pages),
                word_count=len(content.split())
            )
        except ImportError:
            logger.error("pypdf non installé. Installer avec: pip install pypdf")
            raise
        except Exception as e:
            logger.error(f"Erreur parsing PDF: {e}", exc_info=True)
            raise

    async def _parse_docx(self, file_path: str, metadata: Dict) -> ParsedDocument:
        """Parse un fichier DOCX"""
        try:
            from docx import Document

            doc = Document(file_path)

            # Extraire le texte de tous les paragraphes
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            content = "\n\n".join(paragraphs)

            # Extraire métadonnées DOCX
            core_properties = doc.core_properties
            metadata.update({
                "title": core_properties.title or "",
                "author": core_properties.author or "",
                "subject": core_properties.subject or "",
                "keywords": core_properties.keywords or "",
            })

            return ParsedDocument(
                content=content,
                format=DocumentFormat.DOCX,
                metadata=metadata,
                word_count=len(content.split())
            )
        except ImportError:
            logger.error("python-docx non installé. Installer avec: pip install python-docx")
            raise
        except Exception as e:
            logger.error(f"Erreur parsing DOCX: {e}", exc_info=True)
            raise

    async def _parse_text(self, file_path: str, metadata: Dict) -> ParsedDocument:
        """Parse un fichier texte brut"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()

            return ParsedDocument(
                content=content,
                format=DocumentFormat.TXT,
                metadata=metadata,
                word_count=len(content.split())
            )
        except Exception as e:
            logger.error(f"Erreur parsing TXT: {e}", exc_info=True)
            raise

    async def _parse_markdown(self, file_path: str, metadata: Dict) -> ParsedDocument:
        """Parse un fichier Markdown"""
        try:
            import markdown
            from bs4 import BeautifulSoup

            with open(file_path, 'r', encoding='utf-8') as f:
                md_content = f.read()

            # Convertir Markdown en HTML puis extraire le texte
            html = markdown.markdown(md_content)
            soup = BeautifulSoup(html, 'html.parser')
            content = soup.get_text(separator='\n\n')

            # Extraire titre depuis premier heading H1
            first_h1 = soup.find('h1')
            if first_h1 and not metadata.get("title"):
                metadata["title"] = first_h1.get_text()

            return ParsedDocument(
                content=content,
                format=DocumentFormat.MARKDOWN,
                metadata=metadata,
                word_count=len(content.split())
            )
        except ImportError:
            logger.error("markdown/beautifulsoup4 non installés")
            raise
        except Exception as e:
            logger.error(f"Erreur parsing Markdown: {e}", exc_info=True)
            raise

    async def _parse_html(self, file_path: str, metadata: Dict) -> ParsedDocument:
        """Parse un fichier HTML"""
        try:
            from bs4 import BeautifulSoup

            with open(file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()

            soup = BeautifulSoup(html_content, 'html.parser')

            # Extraire métadonnées HTML
            title_tag = soup.find('title')
            if title_tag and not metadata.get("title"):
                metadata["title"] = title_tag.get_text()

            meta_description = soup.find('meta', attrs={'name': 'description'})
            if meta_description and not metadata.get("description"):
                metadata["description"] = meta_description.get('content', '')

            # Retirer scripts et styles
            for script in soup(["script", "style"]):
                script.decompose()

            # Extraire le texte
            content = soup.get_text(separator='\n\n')

            # Nettoyer les lignes vides multiples
            lines = [line.strip() for line in content.splitlines() if line.strip()]
            content = '\n\n'.join(lines)

            return ParsedDocument(
                content=content,
                format=DocumentFormat.HTML,
                metadata=metadata,
                word_count=len(content.split())
            )
        except ImportError:
            logger.error("beautifulsoup4 non installé")
            raise
        except Exception as e:
            logger.error(f"Erreur parsing HTML: {e}", exc_info=True)
            raise

    def get_supported_extensions(self) -> List[str]:
        """Retourne la liste des extensions supportées"""
        return list(self.supported_formats.keys())
